"""Build .docx cover letters (FR + EN) from Groq-generated paragraphs.

Uses python-docx. Replicates the layout of Hugo's Laplace template:
  - Sender header (left, 5 lines, small)
  - Blank line
  - Recipient block (right-indented): company + "Adresse du recruteur" + date
  - Objet line (bold)
  - Salutation
  - 4 body paragraphs, justified
  - Closing + signature
"""
from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


def _fr_date(d: datetime) -> str:
    months = [
        "janvier", "février", "mars", "avril", "mai", "juin",
        "juillet", "août", "septembre", "octobre", "novembre", "décembre",
    ]
    return f"Le {d.day} {months[d.month - 1]} {d.year}"


def _en_date(d: datetime) -> str:
    months = [
        "January", "February", "March", "April", "May", "June",
        "July", "August", "September", "October", "November", "December",
    ]
    day = d.day
    if 10 <= day % 100 <= 20:
        suffix = "th"
    else:
        suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
    return f"{months[d.month - 1]} the {day}{suffix}, {d.year}"


def _slugify(s: str) -> str:
    s = re.sub(r"[^\w\s-]", "", s, flags=re.UNICODE).strip().lower()
    s = re.sub(r"[-\s]+", "-", s)
    return s[:60]


def _build_letter(
    path: Path,
    *,
    lang: str,  # "fr" or "en"
    sender: dict,
    company: str,
    location: str,
    subject: str,
    paragraphs: list[str],
) -> None:
    doc = Document()

    # Page margins 2 cm
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def add_para(text: str = "", *, align=None, bold=False, size=None, space_after=4):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        if text:
            run = p.add_run(text)
            run.bold = bold
            if size:
                run.font.size = Pt(size)
        return p

    # --- Sender block (left) ---
    add_para(sender["name"], bold=True, size=11)
    add_para(sender["address_line1"], size=10)
    add_para(sender["address_line2"], size=10)
    add_para(sender["email"], size=10)
    add_para(sender["phone"], size=10, space_after=14)

    # --- Recipient block (right-indented via paragraph indent) ---
    today = datetime.now()
    date_str = _fr_date(today) if lang == "fr" else _en_date(today)
    recipient_addr = (
        "Adresse du recruteur" if lang == "fr" else "Recruiter address"
    )

    def add_recipient(text: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.size = Pt(10)

    add_recipient(company)
    add_recipient(recipient_addr)
    add_recipient(location or "")
    add_recipient(date_str)

    # Spacer
    add_para("", space_after=10)

    # --- Subject line ---
    objet_label = "Objet" if lang == "fr" else "Subject"
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r1 = p.add_run(f"{objet_label} : ")
    r1.bold = True
    r2 = p.add_run(subject)
    r2.bold = True

    # --- Salutation ---
    salutation = "Madame, Monsieur," if lang == "fr" else "Dear Sir or Madam,"
    add_para(salutation, space_after=12)

    # --- Body paragraphs ---
    for para in paragraphs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.first_line_indent = Cm(0)
        p.add_run(para.strip())

    # --- Closing ---
    add_para("", space_after=6)
    closing = "Cordialement," if lang == "fr" else "Kind regards,"
    add_para(closing, space_after=16)
    add_para(sender["name"], bold=True)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def write_cover_letters(
    output_dir: Path,
    job_id: str,
    title: str,
    company: str,
    location: str,
    sender: dict,
    content: dict,
) -> tuple[Path, Path]:
    """Write both FR and EN .docx files. Returns (fr_path, en_path)."""
    base = f"{_slugify(company)}_{_slugify(title)}_{job_id[:8]}"
    fr_path = output_dir / f"{base}_FR.docx"
    en_path = output_dir / f"{base}_EN.docx"

    _build_letter(
        fr_path,
        lang="fr",
        sender=sender,
        company=company,
        location=location,
        subject=content["fr_subject"],
        paragraphs=content["fr_paragraphs"],
    )
    _build_letter(
        en_path,
        lang="en",
        sender=sender,
        company=company,
        location=location,
        subject=content["en_subject"],
        paragraphs=content["en_paragraphs"],
    )
    return fr_path, en_path
