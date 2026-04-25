"""Cover letter generator — Kairo.

Generate a personalised cover letter for a given job offer, using Groq (same
model family as the rest of the app for consistency). Can output in French or
English. Optionally takes the user's existing cover letter as a style reference.

Public API:
    - generate_cover_letter(cv_text, job, language, user_template) -> str | None
    - build_cover_letter_docx(text, candidate, job) -> bytes

The generated text follows a 4-paragraph structure:
    1. Accroche       (why this role / this company)
    2. Valeur ajoutée (match CV ↔ poste avec 2-3 exemples concrets)
    3. Motivation     (ce qui te tire vers l'entreprise spécifiquement)
    4. Appel à l'action

The DOCX renderer matches the classic French business-letter layout:
    - Candidate block (left-aligned, top)         : name (bold) + email
    - Recipient block (right-aligned)             : company + location
    - Date (right-aligned)                        : "Le 20 novembre 2025" / "November the 20th, 2025"
    - Subject (bold, justified)
    - Body (4 paragraphs, justified)
    - Signature name (right-aligned, below body)
"""
from __future__ import annotations

import io
import json
import os
import time
import urllib.error
import urllib.request
from datetime import date
from typing import Any

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor

GROQ_MODEL = os.environ.get("GROQ_MODEL", "llama-3.3-70b-versatile")
GROQ_URL = "https://api.groq.com/openai/v1/chat/completions"

SUPPORTED_LANGUAGES = ("fr", "en")


# ---------------------------------------------------------------------------
# LLM prompts — one per language, tuned for the French/English job market.
# NB: We deliberately ask the LLM NOT to write the signature name at the end,
#     because the DOCX renderer places it in its own right-aligned paragraph.
# ---------------------------------------------------------------------------
_SYSTEM_FR = """Tu es un coach de carrière senior qui rédige des lettres de motivation percutantes pour le marché français et européen.

Règles STRICTES :
1. Langue : français irréprochable, registre professionnel mais pas pompeux. Tutoiement interdit.
2. Longueur : 280-360 mots. Ni plus, ni moins. Une page A4 bien remplie.
3. Structure en 4 paragraphes clairement distincts (sépare-les par des retours à la ligne simples) :
   - §1 Accroche : pourquoi ce poste dans cette entreprise, 2-3 phrases maximum.
   - §2 Valeur ajoutée : 2 ou 3 exemples concrets tirés du CV qui matchent l'offre.
   - §3 Motivation : ce qui te tire spécifiquement vers cette entreprise (secteur, culture, produit).
   - §4 Appel à l'action : disponibilité pour un échange.
4. Jamais de phrases bateau type "passionné par", "dynamique et motivé", "ma candidature retiendra toute votre attention", "ayant pris connaissance de votre annonce".
5. Pas de liste à puces, pas de titres, pas de markdown : prose continue uniquement.
6. Commence directement par "Madame, Monsieur," (ou "Madame [Nom]," si le recruteur est connu).
7. Termine par une dernière ligne "Cordialement," SEULE. N'écris PAS le nom du candidat à la fin, il sera ajouté séparément au format signature.
8. N'invente JAMAIS d'expériences qui ne sont pas dans le CV. Si une compétence demandée n'apparaît pas, ne la mentionne pas.
9. Si le candidat a une lettre de motivation de référence fournie, inspire-toi de son STYLE (formulations, ton) mais pas de son CONTENU (il sera différent pour chaque poste).

Réponds UNIQUEMENT avec le texte de la lettre (greeting + 4 paragraphes + "Cordialement,"), pas de préambule, pas d'objet, pas de markdown, pas de signature."""

_SYSTEM_EN = """You are a senior career coach writing high-impact cover letters for the UK, US and international job market.

STRICT rules:
1. Language: flawless English, professional but conversational — not stiff.
2. Length: 280-360 words. Exactly one page A4.
3. Four clearly distinct paragraphs (single newline between each):
   - §1 Hook: why this role at this company, 2-3 sentences max.
   - §2 Value: 2 or 3 concrete examples from the CV that directly match the role.
   - §3 Motivation: what specifically draws you to this company (sector, culture, product).
   - §4 Call to action: availability for a chat.
4. NEVER use tired phrases like "I am writing to apply", "detail-oriented team player", "proven track record", "I am confident that", "please find attached".
5. No bullet points, no headers, no markdown: continuous prose only.
6. Start with "Dear Hiring Manager," (or "Dear [Name]," if known).
7. End with a final line "Kind regards," ALONE. Do NOT write the candidate's name at the end — it will be added separately as a signature.
8. NEVER invent experience that isn't in the CV. If a required skill is missing, don't mention it.
9. If the candidate provided a reference cover letter, mirror its STYLE (phrasing, tone) — never its CONTENT (it's for a different role).

Reply ONLY with the letter body text (greeting + 4 paragraphs + "Kind regards,"). No preamble, no subject line, no markdown, no signature."""


# ---------------------------------------------------------------------------
# Groq call (same retry pattern as config_extractor.py)
# ---------------------------------------------------------------------------
def _http_call_groq(api_key: str, payload: dict) -> dict | None:
    data = json.dumps(payload).encode("utf-8")
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}",
        "User-Agent": "kairo-cover-letter/1.0",
    }
    for attempt in range(4):
        try:
            req = urllib.request.Request(GROQ_URL, data=data, headers=headers, method="POST")
            with urllib.request.urlopen(req, timeout=60) as resp:
                return json.loads(resp.read().decode())
        except urllib.error.HTTPError as e:
            if e.code == 429 and attempt < 3:
                wait = 5 * (2 ** attempt)
                time.sleep(wait)
                continue
            try:
                err_body = e.read().decode()[:300]
            except Exception:  # noqa: BLE001
                err_body = ""
            print(f"[cover_letter] HTTP {e.code}: {err_body}")
            return None
        except Exception as e:  # noqa: BLE001
            print(f"[cover_letter] error: {e}")
            return None
    return None


# ---------------------------------------------------------------------------
# Public: generate_cover_letter
# ---------------------------------------------------------------------------
def generate_cover_letter(
    *,
    cv_text: str,
    job: dict[str, Any],
    candidate_name: str,
    language: str = "fr",
    user_template: str | None = None,
    max_retries: int = 2,
) -> str | None:
    """Generate a tailored cover letter.

    Args:
        cv_text:        Raw text of the candidate's CV.
        job:            Dict with keys at least: title, company, description, location (optional).
        candidate_name: Full name used to sign the letter.
        language:       'fr' | 'en'
        user_template:  Optional: the candidate's own existing cover letter, used
                        as a style reference (not as content).
        max_retries:    Extra Groq attempts on transient failure.

    Returns:
        The letter as a plain-text string (greeting + body + sign-off), or None on failure.
    """
    language = (language or "fr").lower()
    if language not in SUPPORTED_LANGUAGES:
        language = "fr"

    api_key = os.environ.get("GROQ_API_KEY") or os.environ.get("GEMINI_API_KEY")
    if not api_key:
        print("[cover_letter] Missing GROQ_API_KEY env var")
        return None

    system = _SYSTEM_FR if language == "fr" else _SYSTEM_EN

    # Truncate inputs to stay in token budget.
    cv_slice   = (cv_text or "")[:7000]
    desc_slice = (job.get("description") or "")[:4500]
    tmpl_slice = (user_template or "")[:2500]

    lines_fr = [
        "=== CANDIDAT ===",
        f"Nom : {candidate_name}",
        "",
        "=== CV ===",
        cv_slice or "(aucun CV)",
        "",
        "=== OFFRE ===",
        f"Poste     : {job.get('title') or '?'}",
        f"Entreprise: {job.get('company') or '?'}",
        f"Lieu      : {job.get('location') or '?'}",
        "",
        "Description :",
        desc_slice or "(description non fournie)",
    ]
    lines_en = [
        "=== CANDIDATE ===",
        f"Name: {candidate_name}",
        "",
        "=== CV ===",
        cv_slice or "(no CV)",
        "",
        "=== JOB ===",
        f"Role    : {job.get('title') or '?'}",
        f"Company : {job.get('company') or '?'}",
        f"Location: {job.get('location') or '?'}",
        "",
        "Description:",
        desc_slice or "(no description)",
    ]
    user_msg = "\n".join(lines_fr if language == "fr" else lines_en)

    if tmpl_slice:
        header = ("\n\n=== LETTRE DE RÉFÉRENCE (STYLE UNIQUEMENT) ===\n"
                  if language == "fr"
                  else "\n\n=== REFERENCE LETTER (STYLE ONLY) ===\n")
        user_msg += header + tmpl_slice

    payload = {
        "model": GROQ_MODEL,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user_msg},
        ],
        "temperature": 0.55,
        "top_p": 0.9,
        "max_tokens": 900,
    }

    for _ in range(max_retries + 1):
        body = _http_call_groq(api_key, payload)
        if not body:
            continue
        try:
            text = body["choices"][0]["message"]["content"].strip()
        except (KeyError, IndexError, TypeError):
            continue
        if text:
            return _sanitize(text, candidate_name=candidate_name, language=language)
    return None


def _sanitize(text: str, *, candidate_name: str, language: str) -> str:
    """Strip accidental markdown fences / whitespace / stray signature lines."""
    lines = [ln.rstrip() for ln in text.splitlines()]
    # Drop markdown code fences if any.
    lines = [ln for ln in lines if not ln.strip().startswith("```")]
    # Collapse triple blank lines into double.
    cleaned: list[str] = []
    blank_run = 0
    for ln in lines:
        if ln.strip() == "":
            blank_run += 1
            if blank_run <= 1:
                cleaned.append("")
        else:
            blank_run = 0
            cleaned.append(ln)

    # Safety net: LLMs sometimes slip the candidate name back after the sign-off.
    # We strip any trailing line that looks like the signature name, so the
    # renderer keeps a single right-aligned signature paragraph.
    name_token = (candidate_name or "").strip().lower()
    while cleaned and cleaned[-1].strip() == "":
        cleaned.pop()
    if name_token and cleaned:
        last_clean = cleaned[-1].strip().lower()
        if last_clean == name_token or last_clean.startswith(name_token):
            cleaned.pop()
            while cleaned and cleaned[-1].strip() == "":
                cleaned.pop()

    return "\n".join(cleaned).strip()


# ---------------------------------------------------------------------------
# Date helpers — match the French business layout style
# ---------------------------------------------------------------------------
_MONTHS_FR = [
    "janvier", "février", "mars", "avril", "mai", "juin",
    "juillet", "août", "septembre", "octobre", "novembre", "décembre",
]


def _format_date_fr(d: date) -> str:
    return f"Le {d.day} {_MONTHS_FR[d.month - 1]} {d.year}"


def _format_date_en(d: date) -> str:
    day = d.day
    if 10 <= day % 100 <= 20:
        suffix = "th"
    else:
        suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
    return f"{d.strftime('%B')} the {day}{suffix}, {d.year}"


def _subject_line(language: str, job_title: str) -> str:
    if language == "fr":
        return f"Objet : Candidature au poste de {job_title}"
    return f"Subject: Application for the position of {job_title}"


# ---------------------------------------------------------------------------
# DOCX renderer — polished A4 letter, layout matching the Laplace template
# ---------------------------------------------------------------------------
def build_cover_letter_docx(
    *,
    text: str,
    candidate_name: str,
    candidate_email: str | None = None,
    job_title: str | None = None,
    job_company: str | None = None,
    job_location: str | None = None,
    language: str = "fr",
) -> bytes:
    """Render the cover letter into a proper A4 Word document.

    Layout (matches the template Hugo provided):
        Candidate name (bold, left)
        Candidate email (small, slate, left)
        (blank)
        Company name (right)
        Company location (right)
        (blank)
        Date (right)
        (blank)
        Subject line (bold, justified)
        (blank)
        Body (justified, 11pt Calibri, 1.25 line spacing)
        (blank)
        Candidate signature name (right)
    """
    language = (language or "fr").lower()
    if language not in SUPPORTED_LANGUAGES:
        language = "fr"

    doc = Document()

    # --- Page & base typography ---
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    base_style = doc.styles["Normal"]
    base_style.font.name = "Calibri"
    base_style.font.size = Pt(11)

    muted = RGBColor(0x64, 0x74, 0x8B)  # slate-500

    def _add_paragraph(
        text_: str | None = None,
        *,
        align: int = WD_PARAGRAPH_ALIGNMENT.LEFT,
        bold: bool = False,
        italic: bool = False,
        size: float | None = None,
        color: RGBColor | None = None,
        space_after: float = 2.0,
        line_spacing: float = 1.25,
    ) -> None:
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after  = Pt(space_after)
        p.paragraph_format.line_spacing = line_spacing
        if text_:
            r = p.add_run(text_)
            if bold:
                r.bold = True
            if italic:
                r.italic = True
            if size is not None:
                r.font.size = Pt(size)
            if color is not None:
                r.font.color.rgb = color

    # --- Candidate header (left) ---
    _add_paragraph(candidate_name, bold=True, size=11.5, space_after=0)
    if candidate_email:
        _add_paragraph(candidate_email, size=10, color=muted, space_after=2)

    _add_paragraph(space_after=2)  # spacer between candidate and recipient

    # --- Recipient block (right) ---
    if job_company:
        _add_paragraph(
            job_company, align=WD_PARAGRAPH_ALIGNMENT.RIGHT, bold=True, space_after=0,
        )
    if job_location:
        _add_paragraph(
            job_location, align=WD_PARAGRAPH_ALIGNMENT.RIGHT, space_after=2,
        )

    _add_paragraph(space_after=2)  # spacer before date

    # --- Date (right) ---
    today = date.today()
    date_str = _format_date_fr(today) if language == "fr" else _format_date_en(today)
    _add_paragraph(date_str, align=WD_PARAGRAPH_ALIGNMENT.RIGHT, space_after=8)

    # --- Subject line (bold, justified) ---
    if job_title:
        subj = doc.add_paragraph()
        subj.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        subj.paragraph_format.space_after = Pt(10)
        r = subj.add_run(_subject_line(language, job_title))
        r.bold = True

    # --- Body (justified) ---
    paragraphs = [blk.strip("\n") for blk in text.split("\n\n") if blk.strip()]
    for i, raw in enumerate(paragraphs):
        p = doc.add_paragraph(raw)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after  = Pt(6)
        p.paragraph_format.line_spacing = 1.25

    # --- Signature (right) ---
    _add_paragraph(space_after=2)
    _add_paragraph(
        candidate_name,
        align=WD_PARAGRAPH_ALIGNMENT.RIGHT,
        bold=True,
        space_after=0,
    )

    # Serialize
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Template mode : clone the user's uploaded .docx as a shell
# ---------------------------------------------------------------------------
#
# The idea : the user uploaded his own cover-letter template in onboarding.
# We want to preserve its exact layout (fonts, margins, header/footer,
# candidate block, signature block) and only swap the *dynamic* parts with
# the generated text and per-offer context.
#
# Primary mechanism : MARKERS in the template. The user can add any of these
# tokens anywhere in his .docx :
#     {{BODY}}         → the generated letter body (multi-paragraph)
#     {{SUBJECT}}      → "Objet : Candidature au poste de X" (FR/EN)
#     {{DATE}}         → "Le 20 novembre 2025" / "November the 20th, 2025"
#     {{COMPANY}}      → recipient company name
#     {{LOCATION}}     → recipient location
#     {{JOB_TITLE}}    → the job title verbatim
#     {{CANDIDATE_NAME}}
#
# Fallback : if no {{BODY}} marker is found, we replace the *largest
# contiguous block of text paragraphs* (i.e. the first run of 2+ paragraphs
# longer than 80 chars each) with the generated body. This works for
# templates that already have "lorem ipsum"-style filler text for the letter.
# ---------------------------------------------------------------------------
_TEMPLATE_BODY_MARKER = "{{BODY}}"


def render_from_template(
    template_bytes: bytes,
    *,
    body_text: str,
    candidate_name: str,
    candidate_email: str | None = None,
    job_title: str | None = None,
    job_company: str | None = None,
    job_location: str | None = None,
    language: str = "fr",
) -> bytes:
    """Clone the user's .docx template, inject dynamic content, return bytes.

    See module docstring for marker semantics.
    """
    language = (language or "fr").lower()
    if language not in SUPPORTED_LANGUAGES:
        language = "fr"

    doc = Document(io.BytesIO(template_bytes))

    today = date.today()
    date_str = _format_date_fr(today) if language == "fr" else _format_date_en(today)
    subject = _subject_line(language, job_title or "") if job_title else ""

    tokens = {
        "{{DATE}}":           date_str,
        "{{COMPANY}}":        job_company or "",
        "{{LOCATION}}":       job_location or "",
        "{{JOB_TITLE}}":      job_title or "",
        "{{SUBJECT}}":        subject,
        "{{CANDIDATE_NAME}}": candidate_name or "",
        "{{CANDIDATE_EMAIL}}": candidate_email or "",
    }

    body_paragraphs = [blk.strip() for blk in (body_text or "").split("\n\n") if blk.strip()]

    body_marker_found = False

    # Walk every paragraph (including those inside tables, if the user uses a
    # table-based layout — the Laplace template does not, but some do).
    for para in _iter_all_paragraphs(doc):
        text = para.text
        if not text:
            continue

        # 1) Body marker — swap this paragraph for N body paragraphs.
        if _TEMPLATE_BODY_MARKER in text:
            _inject_body_at_marker(para, body_paragraphs)
            body_marker_found = True
            continue

        # 2) Simple token replacement — preserves runs as much as possible.
        for token, value in tokens.items():
            if token in para.text:
                _replace_token_in_paragraph(para, token, value)

    # 3) Fallback : no {{BODY}} marker → find the biggest block of text and
    # replace it.
    if not body_marker_found and body_paragraphs:
        _replace_body_region_heuristic(doc, body_paragraphs)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _iter_all_paragraphs(doc):
    """Yield every Paragraph in the document (body + table cells)."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _replace_token_in_paragraph(para, token: str, value: str) -> None:
    """Replace `token` with `value` in `para`, preserving run formatting when
    the whole token lives inside a single run. Otherwise we collapse the
    paragraph's runs into one (minimal formatting loss — usually fine since
    placeholders like {{DATE}} rarely get split across runs)."""
    # Happy path: token confined to one run.
    for run in para.runs:
        if token in run.text:
            run.text = run.text.replace(token, value)
            return

    # Split path: token spans multiple runs (very rare, mostly from copy-paste
    # in Word which sometimes wraps styling around braces). Collapse.
    if token in para.text:
        new_text = para.text.replace(token, value)
        # Keep first run's formatting; blank out the rest.
        if para.runs:
            first = para.runs[0]
            first.text = new_text
            for r in para.runs[1:]:
                r.text = ""
        else:
            para.add_run(new_text)


def _inject_body_at_marker(para, body_paragraphs: list[str]) -> None:
    """Replace a paragraph whose text contains {{BODY}} with N paragraphs
    carrying the generated body, each inheriting the marker paragraph's
    style + formatting.

    Strategy: clone the marker paragraph's XML once per body line, change the
    text, insert the clones BEFORE the marker, then delete the marker. Copying
    the XML preserves indentation, font, alignment, line-spacing etc.
    """
    if not body_paragraphs:
        # No body to inject — just blank the marker so we don't leave "{{BODY}}".
        _replace_token_in_paragraph(para, _TEMPLATE_BODY_MARKER, "")
        return

    from copy import deepcopy
    from docx.oxml.ns import qn

    marker_el = para._element
    parent = marker_el.getparent()
    if parent is None:
        return
    idx = list(parent).index(marker_el)

    for offset, body_line in enumerate(body_paragraphs):
        clone = deepcopy(marker_el)
        # Remove every run inside the clone — we'll add a single clean run.
        for r in clone.findall(qn("w:r")):
            clone.remove(r)
        # Rebuild with one run carrying the body line.
        new_run = _make_run_element(body_line, _pick_template_run_properties(marker_el))
        clone.append(new_run)
        parent.insert(idx + offset, clone)

    # Remove the original marker paragraph.
    parent.remove(marker_el)


def _pick_template_run_properties(paragraph_el):
    """Return the <w:rPr> element of the first run in the paragraph, so we
    can reuse the template's font/size/color for the injected body lines.
    May return None if the paragraph has no explicit run formatting — the
    paragraph-level style will then apply."""
    from docx.oxml.ns import qn
    from copy import deepcopy

    first_run = paragraph_el.find(qn("w:r"))
    if first_run is None:
        return None
    rpr = first_run.find(qn("w:rPr"))
    return deepcopy(rpr) if rpr is not None else None


def _make_run_element(text: str, rpr_template):
    """Build a <w:r><w:t>text</w:t></w:r> element, optionally with the
    provided run-properties <w:rPr>. Adds xml:space="preserve" so leading /
    trailing whitespace is not stripped by Word."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from copy import deepcopy

    r = OxmlElement("w:r")
    if rpr_template is not None:
        r.append(deepcopy(rpr_template))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _replace_body_region_heuristic(doc, body_paragraphs: list[str]) -> None:
    """Fallback when no {{BODY}} marker is present.

    We look for the longest contiguous block of "body-ish" paragraphs —
    heuristic : >= 2 consecutive paragraphs each >= 80 characters — and
    replace that region with the generated paragraphs, preserving the first
    block paragraph's style.
    """
    paras = list(doc.paragraphs)
    if not paras:
        return

    MIN_CHARS = 80
    MIN_RUN = 2

    best_start, best_end, best_len = -1, -1, 0
    run_start = None
    for i, p in enumerate(paras):
        txt = p.text.strip()
        is_body = len(txt) >= MIN_CHARS
        if is_body:
            if run_start is None:
                run_start = i
        else:
            if run_start is not None:
                run_len = i - run_start
                if run_len >= MIN_RUN and run_len > best_len:
                    best_start, best_end, best_len = run_start, i - 1, run_len
                run_start = None
    # Tail run.
    if run_start is not None:
        run_len = len(paras) - run_start
        if run_len >= MIN_RUN and run_len > best_len:
            best_start, best_end, best_len = run_start, len(paras) - 1, run_len

    if best_start < 0:
        # No suitable region found — append body at the end, simple fallback.
        for line in body_paragraphs:
            doc.add_paragraph(line)
        return

    # Reuse the first body paragraph as the "anchor" — clone it for every
    # new body line, then remove the old region.
    anchor = paras[best_start]._element
    parent = anchor.getparent()
    if parent is None:
        return
    insert_idx = list(parent).index(anchor)

    from copy import deepcopy
    from docx.oxml.ns import qn

    clones = []
    for line in body_paragraphs:
        clone = deepcopy(anchor)
        for r in clone.findall(qn("w:r")):
            clone.remove(r)
        clone.append(_make_run_element(line, _pick_template_run_properties(anchor)))
        clones.append(clone)

    # Remove the whole region (inclusive).
    for p in paras[best_start: best_end + 1]:
        parent.remove(p._element)

    for offset, clone in enumerate(clones):
        parent.insert(insert_idx + offset, clone)
