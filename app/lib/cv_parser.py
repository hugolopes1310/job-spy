"""CV / cover letter parsing — PDF and DOCX to plain text.

Intentionally simple: we trust the LLM downstream to understand structure,
so we just need clean-enough text (no section detection, no NER).

Error model — every distinct failure mode raises a specific subclass of
:class:`CVParseError` so the UI can show a friendly, actionable message
instead of a stack trace.
"""
from __future__ import annotations

import io
import re
from pathlib import Path


# ---------------------------------------------------------------------------
# Error hierarchy
# ---------------------------------------------------------------------------
class CVParseError(ValueError):
    """Base class — parsing failed for a known reason. UI-safe message."""


class UnsupportedFormatError(CVParseError):
    """Raised when the uploaded file extension is not .pdf / .docx / .txt."""


class FileTooLargeError(CVParseError):
    """Uploaded file is larger than the safety cap."""


class EmptyFileError(CVParseError):
    """Zero-byte upload."""


class CorruptedFileError(CVParseError):
    """File extension is supported but the binary blob is malformed."""


class EncryptedPDFError(CVParseError):
    """Password-protected PDF — pypdf can open it but extract_text yields nothing."""


class ImageOnlyPDFError(CVParseError):
    """PDF that contains only images / scans — no extractable text."""


class CVParseConfigError(RuntimeError):
    """Missing dependency (pypdf / python-docx). Surfaces as a config error,
    not a CV parse error, because it's a deploy-time issue, not user input."""


# ---------------------------------------------------------------------------
# Limits — keep them generous but bounded.
# ---------------------------------------------------------------------------
MAX_FILE_BYTES = 10 * 1024 * 1024  # 10 MiB
# Below this, we suspect an image-only PDF or a scan with bad OCR.
MIN_TEXT_CHARS_FOR_PDF = 80


def _clean(text: str) -> str:
    """Collapse whitespace, strip zero-width chars, normalize line breaks."""
    if not text:
        return ""
    # Remove common zero-width / bidi chars that mess up LLM tokenization.
    text = re.sub(r"[\u200b-\u200f\u202a-\u202e\ufeff]", "", text)
    # Normalize line endings.
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Collapse 3+ newlines to 2 (preserve paragraph breaks).
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Collapse horizontal whitespace.
    text = re.sub(r"[ \t]+", " ", text)
    # Trim each line.
    text = "\n".join(line.strip() for line in text.split("\n"))
    return text.strip()


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------
def _parse_pdf(file_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader
        from pypdf.errors import PdfReadError  # type: ignore[import]
    except ImportError as e:
        raise CVParseConfigError(
            "pypdf manquant — `pip install pypdf` (déjà dans requirements.txt)"
        ) from e

    try:
        reader = PdfReader(io.BytesIO(file_bytes))
    except PdfReadError as e:
        raise CorruptedFileError(
            "PDF illisible (fichier corrompu ou tronqué). Réessaie avec un autre PDF."
        ) from e
    except Exception as e:  # noqa: BLE001
        raise CorruptedFileError(
            f"Impossible d'ouvrir le PDF : {type(e).__name__}."
        ) from e

    # Encrypted PDFs may "open" but yield no text — try to detect both cases.
    if getattr(reader, "is_encrypted", False):
        # Best-effort: try empty-password decrypt before giving up.
        try:
            ok = reader.decrypt("")
        except Exception:  # noqa: BLE001
            ok = 0
        if not ok:
            raise EncryptedPDFError(
                "Ce PDF est protégé par mot de passe. Enlève la protection puis réessaie."
            )

    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001
            # Don't crash on one bad page — skip it.
            continue
    text = "\n\n".join(parts)
    # If the text content is suspiciously short, we likely have an image-only
    # PDF (scan, photo). Tell the user what to do instead of returning garbage.
    if len(text.strip()) < MIN_TEXT_CHARS_FOR_PDF:
        raise ImageOnlyPDFError(
            "Ce PDF semble être une image scannée (texte non extractible). "
            "Fournis plutôt un PDF avec du texte sélectionnable, ou un .docx."
        )
    return text


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------
def _parse_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError  # type: ignore[import]
    except ImportError as e:
        raise CVParseConfigError(
            "python-docx manquant — `pip install python-docx`"
        ) from e

    try:
        doc = Document(io.BytesIO(file_bytes))
    except PackageNotFoundError as e:
        # Most common: user uploaded a .doc (binary old Word) renamed to .docx,
        # or the file is corrupted.
        raise CorruptedFileError(
            "DOCX illisible. Vérifie que le fichier n'est pas un ancien .doc renommé "
            "et qu'il n'est pas corrompu."
        ) from e
    except Exception as e:  # noqa: BLE001
        raise CorruptedFileError(
            f"Impossible d'ouvrir le DOCX : {type(e).__name__}."
        ) from e

    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)
    # Also pull text out of tables (CVs often use tables for layout).
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------
def parse_cv(file_bytes: bytes, filename: str) -> str:
    """Extract plain text from an uploaded CV / cover letter.

    Args:
        file_bytes: Raw file content.
        filename: Original name (used to detect extension).

    Returns:
        Cleaned plain text.

    Raises:
        UnsupportedFormatError: Extension isn't .pdf / .docx / .txt / .md.
        EmptyFileError:         Zero-byte upload.
        FileTooLargeError:      File exceeds MAX_FILE_BYTES.
        CorruptedFileError:     Header looks valid but content is unreadable.
        EncryptedPDFError:      PDF needs a password.
        ImageOnlyPDFError:      PDF has no extractable text (scan / image-only).
        CVParseConfigError:     Required parsing library is not installed.
    """
    if not file_bytes:
        raise EmptyFileError("Fichier vide. Vérifie que ton CV n'est pas en cours de synchronisation.")
    if len(file_bytes) > MAX_FILE_BYTES:
        raise FileTooLargeError(
            f"Fichier trop volumineux ({len(file_bytes) / 1_048_576:.1f} Mo). "
            f"Taille max : {MAX_FILE_BYTES // 1_048_576} Mo."
        )

    ext = Path(filename).suffix.lower().lstrip(".")
    if ext == "pdf":
        raw = _parse_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        if ext == "doc":
            # python-docx can't open binary .doc — fail fast with a clear msg
            # rather than the generic PackageNotFoundError.
            raise UnsupportedFormatError(
                "Format .doc (ancien Word) non supporté. Convertis-le en .docx avant l'upload."
            )
        raw = _parse_docx(file_bytes)
    elif ext in ("txt", "md"):
        try:
            raw = file_bytes.decode("utf-8", errors="replace")
        except Exception as e:  # noqa: BLE001
            raise CorruptedFileError(f"Impossible de lire le fichier texte : {e}") from e
    else:
        raise UnsupportedFormatError(
            f"Format .{ext} non supporté. Formats acceptés : PDF, DOCX, TXT, MD."
        )
    return _clean(raw)


def parse_cv_from_path(path: str | Path) -> str:
    """Convenience wrapper for local CLI / smoke tests."""
    p = Path(path)
    return parse_cv(p.read_bytes(), p.name)


if __name__ == "__main__":
    import sys

    if len(sys.argv) != 2:
        print("usage: python -m app.lib.cv_parser <path/to/cv.pdf>")
        sys.exit(1)
    text = parse_cv_from_path(sys.argv[1])
    print(f"--- {len(text)} chars ---")
    print(text[:1500])
    if len(text) > 1500:
        print(f"\n... ({len(text) - 1500} more chars)")
